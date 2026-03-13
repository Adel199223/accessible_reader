from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from typing import Iterable
from urllib.parse import urlsplit

from bs4 import BeautifulSoup
from bs4.element import NavigableString, Tag
from docx import Document as DocxDocument
from markdown import markdown
from pypdf import PdfReader

from .models import ViewBlock
from .text_utils import normalize_whitespace


class UnsupportedDocumentError(ValueError):
    """Raised when a document cannot be meaningfully parsed."""


ORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)(?P<index>\d+)[.)]\s+(?P<text>.+)$")
UNORDERED_LIST_RE = re.compile(r"^(?P<indent>\s*)[-*•]\s+(?P<text>.+)$")
QUOTE_PREFIX_RE = re.compile(r"^(?P<prefix>(?:>\s*)+)(?P<text>.+)$")


@dataclass(slots=True)
class ParsedDocument:
    title: str
    source_type: str
    file_name: str | None
    source_locator: str | None
    blocks: list[ViewBlock]
    plain_text: str


def parse_text_input(text: str, title: str | None = None) -> ParsedDocument:
    blocks = _parse_plain_text_blocks(text)
    if not blocks:
        raise UnsupportedDocumentError("The pasted text did not contain readable content.")
    chosen_title = title or _pick_title_from_blocks(blocks) or "Pasted document"
    return ParsedDocument(
        title=chosen_title,
        source_type="paste",
        file_name=None,
        source_locator=None,
        blocks=blocks,
        plain_text=_blocks_to_text(blocks),
    )


def parse_uploaded_file(file_name: str, content: bytes) -> ParsedDocument:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".txt":
        blocks = _parse_plain_text_blocks(content.decode("utf-8", errors="ignore"))
    elif suffix == ".md":
        blocks = _parse_markdown_blocks(content.decode("utf-8", errors="ignore"))
    elif suffix in {".html", ".htm"}:
        blocks = _parse_html_blocks(content.decode("utf-8", errors="ignore"))
    elif suffix == ".docx":
        blocks = _parse_docx_blocks(content)
    elif suffix == ".pdf":
        blocks = _parse_pdf_blocks(content)
    else:
        raise UnsupportedDocumentError(f"`{suffix}` is not supported yet.")

    if not blocks:
        raise UnsupportedDocumentError("The file did not contain readable text.")

    return ParsedDocument(
        title=_pick_title_from_blocks(blocks) or Path(file_name).stem.replace("_", " ").strip() or "Imported document",
        source_type=suffix.removeprefix("."),
        file_name=file_name,
        source_locator=None,
        blocks=blocks,
        plain_text=_blocks_to_text(blocks),
    )


def parse_web_page(url: str, html: str) -> ParsedDocument:
    soup = BeautifulSoup(html, "html.parser")
    title = _pick_web_title(soup, url)
    container = _pick_web_container(soup)
    if container is None:
        raise UnsupportedDocumentError("The webpage did not contain a readable article.")

    _strip_noise_tags(container)
    blocks = _collect_html_blocks(container)
    if not blocks:
        raise UnsupportedDocumentError("The webpage did not contain a readable article.")

    return ParsedDocument(
        title=title,
        source_type="web",
        file_name=None,
        source_locator=url,
        blocks=blocks,
        plain_text=_blocks_to_text(blocks),
    )


def _parse_plain_text_blocks(text: str) -> list[ViewBlock]:
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
    blocks: list[ViewBlock] = []
    paragraph_lines: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        nonlocal index
        if not paragraph_lines:
            return
        paragraph = normalize_whitespace(" ".join(paragraph_lines))
        if paragraph:
            blocks.append(ViewBlock(id=f"block-{index}", kind="paragraph", text=paragraph))
            index += 1
        paragraph_lines.clear()

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue
        ordered_match = ORDERED_LIST_RE.match(raw_line)
        unordered_match = UNORDERED_LIST_RE.match(raw_line)
        quote_match = QUOTE_PREFIX_RE.match(raw_line.strip())
        if ordered_match:
            flush_paragraph()
            indent = ordered_match.group("indent")
            blocks.append(
                ViewBlock(
                    id=f"block-{index}",
                    kind="list_item",
                    text=normalize_whitespace(ordered_match.group("text")),
                    metadata={
                        "list_depth": len(indent) // 2,
                        "list_index": int(ordered_match.group("index")),
                        "list_ordered": True,
                        "source_tag": "plaintext_list_item",
                    },
                )
            )
            index += 1
            continue
        if unordered_match:
            flush_paragraph()
            indent = unordered_match.group("indent")
            blocks.append(
                ViewBlock(
                    id=f"block-{index}",
                    kind="list_item",
                    text=normalize_whitespace(unordered_match.group("text")),
                    metadata={
                        "list_depth": len(indent) // 2,
                        "list_index": index + 1,
                        "list_ordered": False,
                        "source_tag": "plaintext_list_item",
                    },
                )
            )
            index += 1
            continue
        if quote_match:
            flush_paragraph()
            prefix = quote_match.group("prefix")
            quote_depth = prefix.count(">")
            blocks.append(
                ViewBlock(
                    id=f"block-{index}",
                    kind="quote",
                    text=normalize_whitespace(quote_match.group("text")),
                    metadata={
                        "quote_depth": quote_depth,
                        "source_tag": "plaintext_quote",
                    },
                )
            )
            index += 1
            continue
        if _looks_like_heading(line):
            flush_paragraph()
            blocks.append(ViewBlock(id=f"block-{index}", kind="heading", text=line, level=2))
            index += 1
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    return blocks


def _parse_markdown_blocks(text: str) -> list[ViewBlock]:
    blocks = _parse_html_blocks(markdown(text, extensions=["extra", "sane_lists"]))
    if blocks:
        return blocks
    return _parse_plain_text_blocks(text)


def _parse_html_blocks(text: str) -> list[ViewBlock]:
    soup = BeautifulSoup(text, "html.parser")
    container = soup.body or soup
    _strip_noise_tags(container)
    return _collect_html_blocks(container)


def _parse_docx_blocks(content: bytes) -> list[ViewBlock]:
    document = DocxDocument(BytesIO(content))
    blocks: list[ViewBlock] = []
    index = 0
    for paragraph in document.paragraphs:
        text = normalize_whitespace(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        if style_name.startswith("heading"):
            level = 2
            for value in range(1, 5):
                if str(value) in style_name:
                    level = value
                    break
            blocks.append(ViewBlock(id=f"block-{index}", kind="heading", text=text, level=level))
        elif "list" in style_name or text.startswith(("- ", "* ", "• ")):
            ordered = "number" in style_name
            blocks.append(
                ViewBlock(
                    id=f"block-{index}",
                    kind="list_item",
                    text=text.lstrip("-*• ").strip(),
                    metadata={
                        "list_depth": 0,
                        "list_index": index + 1,
                        "list_ordered": ordered,
                        "source_tag": "docx_list_item",
                    },
                )
            )
        else:
            blocks.append(ViewBlock(id=f"block-{index}", kind="paragraph", text=text))
        index += 1
    return blocks


def _parse_pdf_blocks(content: bytes) -> list[ViewBlock]:
    reader = PdfReader(BytesIO(content))
    page_text: list[str] = []
    for page in reader.pages:
        extracted = normalize_whitespace(page.extract_text() or "")
        if extracted:
            page_text.append(extracted)
    if not page_text:
        raise UnsupportedDocumentError("This PDF looks image-based or scanned, which is not supported yet.")
    combined = "\n\n".join(page_text)
    return _parse_plain_text_blocks(combined)


def _pick_title_from_blocks(blocks: Iterable[ViewBlock]) -> str | None:
    for block in blocks:
        if block.kind == "heading":
            return block.text
    for block in blocks:
        if len(block.text) <= 80:
            return block.text
    return None


def _blocks_to_text(blocks: Iterable[ViewBlock]) -> str:
    return "\n".join(block.text for block in blocks)


def _looks_like_heading(line: str) -> bool:
    if len(line) > 80 or line.endswith((".", "!", "?")):
        return False
    words = line.split()
    if len(words) > 8:
        return False
    capitals = sum(1 for word in words if word[:1].isupper())
    return capitals >= max(1, len(words) // 2)


def _collect_html_blocks(container: Tag | BeautifulSoup) -> list[ViewBlock]:
    blocks: list[ViewBlock] = []
    index = 0

    def append_block(kind: str, text: str, *, level: int | None = None, metadata: dict[str, object] | None = None) -> None:
        nonlocal index
        text_value = normalize_whitespace(text)
        if not text_value:
            return
        blocks.append(
            ViewBlock(
                id=f"block-{index}",
                kind=kind,  # type: ignore[arg-type]
                text=text_value,
                level=level,
                metadata={key: value for key, value in (metadata or {}).items() if value is not None},
            )
        )
        index += 1

    def walk(node: Tag | BeautifulSoup, *, quote_depth: int = 0) -> None:
        for child in node.children:
            if isinstance(child, NavigableString):
                continue
            if child.name in {"ol", "ul"}:
                walk_list(child, list_depth=0, quote_depth=quote_depth)
                continue
            if child.name == "blockquote":
                walk(child, quote_depth=quote_depth + 1)
                continue
            if child.name and child.name.startswith("h"):
                kind = "quote" if quote_depth > 0 else "heading"
                level = None if quote_depth > 0 else min(max(int(child.name[1]), 1), 6)
                append_block(
                    kind,
                    child.get_text(" ", strip=True),
                    level=level,
                    metadata={
                        "quote_depth": quote_depth if quote_depth > 0 else None,
                        "source_tag": child.name,
                    },
                )
                continue
            if child.name == "p":
                append_block(
                    "quote" if quote_depth > 0 else "paragraph",
                    child.get_text(" ", strip=True),
                    metadata={
                        "quote_depth": quote_depth if quote_depth > 0 else None,
                        "source_tag": child.name,
                    },
                )
                continue
            walk(child, quote_depth=quote_depth)

    def walk_list(list_node: Tag, *, list_depth: int, quote_depth: int) -> None:
        start_index = int(list_node.get("start") or 1)
        ordered = list_node.name == "ol"
        for item_offset, list_item in enumerate(list_node.find_all("li", recursive=False)):
            list_index = int(list_item.get("value") or (start_index + item_offset))
            append_block(
                "list_item",
                _collect_list_item_text(list_item),
                metadata={
                    "list_depth": list_depth,
                    "list_index": list_index,
                    "list_ordered": ordered,
                    "quote_depth": quote_depth if quote_depth > 0 else None,
                    "source_tag": "li",
                },
            )
            walk_nested_list_item_children(list_item, list_depth=list_depth, quote_depth=quote_depth)

    def walk_nested_list_item_children(list_item: Tag, *, list_depth: int, quote_depth: int) -> None:
        for child in list_item.children:
            if isinstance(child, NavigableString):
                continue
            if child.name in {"ol", "ul"}:
                walk_list(child, list_depth=list_depth + 1, quote_depth=quote_depth)
                continue
            if child.name == "blockquote":
                walk(child, quote_depth=quote_depth + 1)
                continue
            walk_nested_list_item_children(child, list_depth=list_depth, quote_depth=quote_depth)

    walk(container)
    return blocks


def _collect_list_item_text(list_item: Tag) -> str:
    parts: list[str] = []
    for child in list_item.children:
        if isinstance(child, NavigableString):
            text = normalize_whitespace(str(child))
            if text:
                parts.append(text)
            continue
        if child.name in {"ol", "ul", "blockquote"}:
            continue
        child_text = _collect_text_without_structures(child)
        if child_text:
            parts.append(child_text)
    return normalize_whitespace(" ".join(parts))


def _collect_text_without_structures(node: Tag) -> str:
    parts: list[str] = []
    for child in node.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
            continue
        if child.name in {"ol", "ul", "blockquote"}:
            continue
        parts.append(_collect_text_without_structures(child))
    return normalize_whitespace(" ".join(part for part in parts if part))


def _strip_noise_tags(container: Tag | BeautifulSoup) -> None:
    for tag_name in ["script", "style", "noscript", "nav", "header", "footer", "aside", "form", "dialog"]:
        for tag in container.find_all(tag_name):
            tag.decompose()

    noise_keywords = (
        "cookie",
        "consent",
        "banner",
        "advert",
        "promo",
        "newsletter",
        "subscribe",
        "share",
        "social",
        "recommend",
        "related",
        "comment",
        "breadcrumb",
        "sidebar",
        "popover",
        "modal",
    )
    for tag in list(container.find_all(True)):
        metadata = " ".join(
            [
                tag.get("id") or "",
                " ".join(tag.get("class", [])) if isinstance(tag.get("class"), list) else "",
                tag.get("role") or "",
                tag.get("aria-label") or "",
            ]
        ).lower()
        if metadata and any(keyword in metadata for keyword in noise_keywords):
            tag.decompose()


def _pick_web_title(soup: BeautifulSoup, url: str) -> str:
    for attributes in (
        {"property": "og:title"},
        {"name": "twitter:title"},
        {"property": "twitter:title"},
    ):
        meta = soup.find("meta", attrs=attributes)
        if meta and meta.get("content"):
            title = normalize_whitespace(meta["content"])
            if title:
                return title

    if soup.title and soup.title.string:
        title = normalize_whitespace(soup.title.string)
        if title:
            return title

    first_heading = soup.find(["h1", "h2"])
    if first_heading:
        heading_text = normalize_whitespace(first_heading.get_text(" ", strip=True))
        if heading_text:
            return heading_text

    hostname = urlsplit(url).hostname or "Imported article"
    return hostname.removeprefix("www.")


def _pick_web_container(soup: BeautifulSoup) -> Tag | BeautifulSoup | None:
    for selector in ("article", "main", '[role="main"]'):
        container = soup.select_one(selector)
        if container is not None:
            return container

    body = soup.body
    if body is None:
        return None

    best_container: Tag | None = None
    best_score = 0
    for candidate in body.find_all(["section", "div"], recursive=True):
        paragraph_texts = [
            normalize_whitespace(node.get_text(" ", strip=True))
            for node in candidate.find_all("p")
        ]
        paragraph_texts = [text for text in paragraph_texts if text]
        if len(paragraph_texts) < 2:
            continue

        text_length = sum(len(text) for text in paragraph_texts)
        heading_bonus = 120 if candidate.find(["h1", "h2", "h3"]) else 0
        list_bonus = 30 * len(candidate.find_all("li"))
        link_penalty = sum(
            len(normalize_whitespace(link.get_text(" ", strip=True)))
            for link in candidate.find_all("a")
        )
        score = text_length + heading_bonus + list_bonus - min(link_penalty, text_length // 2)
        if score > best_score:
            best_score = score
            best_container = candidate

    return best_container if best_score >= 240 else None
