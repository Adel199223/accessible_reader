from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from app.parsers import parse_text_input, parse_uploaded_file


def test_parse_text_input_extracts_heading_and_list_items() -> None:
    parsed = parse_text_input(
        "Readable Title\n\nThis is the first paragraph.\n\n- First item\n- Second item",
    )

    assert parsed.title == "Readable Title"
    assert [block.kind for block in parsed.blocks] == ["heading", "paragraph", "list_item", "list_item"]


def test_parse_markdown_file_preserves_structure() -> None:
    parsed = parse_uploaded_file(
        "guide.md",
        b"# Main heading\n\nA paragraph for the reader.\n\n- one\n- two",
    )

    assert parsed.title == "Main heading"
    assert parsed.blocks[0].kind == "heading"
    assert parsed.blocks[2].kind == "list_item"


def test_parse_markdown_file_preserves_ordered_lists_nesting_and_quotes() -> None:
    parsed = parse_uploaded_file(
        "guide.md",
        (
            b"# Main heading\n\n"
            b"1. First ordered item\n"
            b"2. Second ordered item\n"
            b"    - Nested bullet\n\n"
            b"> Quoted line"
        ),
    )

    assert parsed.blocks[0].kind == "heading"
    assert parsed.blocks[0].level == 1

    ordered_items = [
        block
        for block in parsed.blocks
        if block.kind == "list_item" and block.metadata.get("list_ordered") is True
    ]
    assert len(ordered_items) == 2
    assert all(block.metadata.get("list_depth") == 0 for block in ordered_items)

    nested_item = next(
        block
        for block in parsed.blocks
        if block.kind == "list_item" and block.metadata.get("list_ordered") is False
    )
    assert nested_item.metadata["list_depth"] == 1

    quote_block = next(block for block in parsed.blocks if block.kind == "quote")
    assert quote_block.metadata["quote_depth"] == 1


def test_parse_docx_file_extracts_heading() -> None:
    document = DocxDocument()
    document.add_heading("Docx heading", level=1)
    document.add_paragraph("Docx paragraph.")
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_uploaded_file("reader.docx", buffer.getvalue())

    assert parsed.title == "Docx heading"
    assert parsed.blocks[0].kind == "heading"


def test_parse_pdf_file_extracts_text() -> None:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(100, 750, "PDF heading")
    pdf.drawString(100, 730, "This PDF contains readable text.")
    pdf.save()

    parsed = parse_uploaded_file("reader.pdf", buffer.getvalue())

    assert "PDF heading" in parsed.title
    assert len(parsed.blocks) >= 1
